import os
import io
from fastapi import FastAPI, Header, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

SERVICE_SECRET = os.getenv("SERVICE_SECRET")

app = FastAPI()


def verify_secret(x_service_secret: str):
    if not SERVICE_SECRET or x_service_secret != SERVICE_SECRET:
        raise HTTPException(status_code=401, detail="Unauthorized")


class Section(BaseModel):
    title: str
    classification: str  # "pre" | "post" | "unclear"
    summary: str


class GenerateRequest(BaseModel):
    patient_name: str
    dob: str
    dol: str
    sections: List[Section]


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


CLASSIFICATION_LABELS = {
    "pre": "[Pre-Incident]",
    "post": "[Post-Incident]",
    "unclear": "[Unclear]",
}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/generate")
def generate(
    body: GenerateRequest,
    x_service_secret: str = Header(...),
):
    verify_secret(x_service_secret)

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Medical Summary — {body.patient_name}")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(
        f"DOB: {body.dob} | Date of Loss: {body.dol}"
    )
    subtitle_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    for i, section in enumerate(body.sections):
        label = CLASSIFICATION_LABELS.get(section.classification, f"[{section.classification}]")

        # Bold heading with classification label
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(f"{section.title}  {label}")
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        # Summary paragraph
        doc.add_paragraph(section.summary)

        # Horizontal rule between sections (not after last)
        if i < len(body.sections) - 1:
            add_horizontal_rule(doc)

    # Save to bytes buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = body.patient_name.replace(" ", "_")
    filename = f"summary_{safe_name}.docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return StreamingResponse(
        buf,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
